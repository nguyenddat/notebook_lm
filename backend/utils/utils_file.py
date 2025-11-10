import os
from typing import List

from docx import Document as DocxDocument
from pdf2image import convert_from_path
from pytesseract import image_to_string
from langchain.schema import Document
from langchain_community.document_loaders import PyPDFLoader

def load_file(file_path: str) -> List[Document]:
    extension = os.path.splitext(file_path)[1].lower()

    if extension.lower() in [".txt", ".md"]:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text, metadata={"page": 1})]
    
    elif extension in [".doc", ".docx"]:
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return [Document(page_content=text, metadata={"page": 1})]

    elif extension.lower() == ".pdf":
        try:
            print("Trying to load PDF without OCR...")
            loader = PyPDFLoader(file_path)
            docs = loader.load()
            for d in docs:
                if "page" not in d.metadata and "page_number" in d.metadata:
                    d.metadata["page"] = d.metadata["page_number"]
            return docs
        
        except Exception as err:
            print(f"Failed to load PDF without OCR, using OCR now...: {err}")
            pages = convert_from_path(file_path, dpi=300)
            docs = []
            for i, page in enumerate(pages, start=1):
                text = image_to_string(page, lang="vie")
                docs.append(Document(page_content=text, metadata={"page": i}))
            return docs

    else:
        raise ValueError(f"Unsupported file extension: {extension}")

